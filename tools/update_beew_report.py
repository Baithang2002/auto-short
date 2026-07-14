from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SOURCE = Path(r"C:\Users\nicit\Downloads\Beew_Work_Accomplishment_Report_Nishita_1 (1) (1).docx")
OUTPUT = Path(r"C:\Users\nicit\.antigravity-ide\projects\auto-short\docs\Beew_Work_Accomplishment_Report_Nishita_Updated.docx")

ROWS = [
    ("Brand Assets", 4, 12, 16),
    ("Brand Guidelines / Documentation", 2, 3, 5),
    ("Templates", 5, 11, 16),
    ("Instagram Carousels", 45, 8, 53),
    ("Instagram Posts", 5, 1, 6),
    ("Instagram Stories", 8, 2, 10),
    ("Instagram Story Highlight Cover Icons", 10, 2, 12),
    ("LinkedIn Posts", 4, 1, 5),
    ("Reels", 4, 1, 5),
    ("Instagram Reel Covers", 1, 1, 2),
    ("LinkedIn Article Banners", 1, 1, 2),
    ("YouTube Thumbnails", 1, 3, 4),
    ("X (Twitter) Header Banners", 0, 1, 1),
    ("Facebook Covers", 0, 0, 0),
    ("Overall Total", 90, 47, 137),
]

NOTE = (
    "Verified against the Google Drive 'Brand Assets' folder and completed Beew-related Canva designs. "
    "Google Drive counts exported files, Canva counts design records, and multi-page Canva designs are counted as "
    "one design. No completed Facebook cover was found in either source, and the X (Twitter) header banner was "
    "found in Canva only."
)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_run(run, bold=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(9.5)
    run.bold = bold


def style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.08


def build_table(doc, table_style=None):
    table = doc.add_table(rows=1, cols=4)
    if table_style is not None:
        table.style = table_style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = [Inches(3.0), Inches(1.35), Inches(1.35), Inches(1.2)]
    headers = [
        "Category",
        "Google Drive\nCompleted Files",
        "Canva\nCompleted Designs",
        "Combined\nTotal",
    ]

    header_cells = table.rows[0].cells
    for idx, text in enumerate(headers):
        header_cells[idx].width = widths[idx]
        header_cells[idx].text = ""
        p = header_cells[idx].paragraphs[0]
        style_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(text)
        style_run(r, bold=True)
        set_cell_shading(header_cells[idx], "D9E8F5")
        set_cell_margins(header_cells[idx])
        header_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for category, drive_count, canva_count, total in ROWS:
        row_cells = table.add_row().cells
        values = [category, str(drive_count), str(canva_count), str(total)]
        for idx, value in enumerate(values):
            row_cells[idx].width = widths[idx]
            row_cells[idx].text = ""
            align = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p = row_cells[idx].paragraphs[0]
            style_paragraph(p, align)
            r = p.add_run(value)
            style_run(r, bold=(category == "Overall Total"))
            set_cell_margins(row_cells[idx])
            row_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if category == "Overall Total":
                set_cell_shading(row_cells[idx], "EEF4EA")

    return table


def build_note(doc):
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph)
    run = paragraph.add_run(NOTE)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(9)
    run.italic = True
    return paragraph


def main():
    doc = Document(str(SOURCE))

    target_table = None
    for table in doc.tables:
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if header == ["Category", "Assets Included", "Quantity"]:
            target_table = table
            break

    if target_table is None:
        raise RuntimeError("Could not find the existing incomplete assets table.")

    target_style = target_table.style
    new_table = build_table(doc, table_style=target_style)
    note_paragraph = build_note(doc)

    target_table._element.addprevious(new_table._element)
    new_table._element.addnext(note_paragraph._element)
    target_table._element.getparent().remove(target_table._element)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
